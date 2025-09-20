import PyPDF2
import docx
import os
import re
import cv2
import numpy as np
import pytesseract
from typing import Optional, Dict, List
from PIL import Image
import easyocr
import fitz  # PyMuPDF
import tempfile
import unicodedata
from rapidfuzz import process, fuzz

# Initialize OCR reader once
try:
    ocr_reader = easyocr.Reader(['en'], gpu=False)  # load only once
    easyocr_available = True
except Exception as e:
    print(f"⚠️ EasyOCR not available: {e}")
    easyocr_available = False

class ResumeParser:
    """Advanced resume parser with enhanced OCR and text processing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
    
    def preprocess_image_for_ocr(self, image_path, resize_scale=1.5, denoise=True, threshold=True):
        """
        Preprocess image for better OCR results
        Returns a numpy array ready for EasyOCR/Tesseract
        """
        try:
            img = cv2.imread(image_path, cv2.IMREAD_COLOR)
            if img is None:
                raise Exception("cv2.imread failed for " + image_path)

            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

            # Upscale small images
            if resize_scale and resize_scale != 1.0:
                new_w = int(gray.shape[1] * resize_scale)
                new_h = int(gray.shape[0] * resize_scale)
                gray = cv2.resize(gray, (new_w, new_h), interpolation=cv2.INTER_LINEAR)

            # Denoise
            if denoise:
                gray = cv2.fastNlMeansDenoising(gray, None, h=10, templateWindowSize=7, searchWindowSize=21)

            # Deskew (estimate angle and rotate)
            coords = np.column_stack(np.where(gray > 0))
            if coords.size:
                angle = cv2.minAreaRect(coords)[-1]
                if angle < -45:
                    angle = -(90 + angle)
                else:
                    angle = -angle
                if abs(angle) > 0.1:
                    (h, w) = gray.shape[:2]
                    M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
                    gray = cv2.warpAffine(gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)

            # Adaptive threshold for better contrast
            if threshold:
                gray = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                            cv2.THRESH_BINARY, 11, 2)

            # Return as RGB numpy array
            rgb = cv2.cvtColor(gray, cv2.COLOR_GRAY2RGB)
            return rgb
        except Exception as e:
            print(f"Image preprocessing error: {e}")
            # Fallback to original image
            return cv2.imread(image_path)
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from supported resume formats with enhanced OCR"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
    
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        try:
            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
                if not text.strip() or len(text) < 30:  # fallback for scanned PDFs
                    text = self._extract_pdf_ocr(file_path)
                return self._clean_text(text)

            elif file_extension in ['.docx', '.doc']:
                text = self._extract_docx_text(file_path)
                return self._clean_text(text)

            elif file_extension in ['.png', '.jpg', '.jpeg']:
                text = self._extract_image_text(file_path)
                return self._clean_text(text)

            elif file_extension == '.txt':
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                return self._clean_text(text)

        except Exception as e:
            raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                except Exception as e:
                    print(f"Warning: Could not extract PDF page: {e}")
        return text
    
    def _extract_pdf_ocr(self, file_path: str) -> str:
        """Fallback OCR for scanned PDFs"""
        text = ""
        doc = fitz.open(file_path)
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            
            # Save to temporary file for OCR processing
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
                tmp_file.write(img_bytes)
                tmp_path = tmp_file.name
            
            try:
                text += self._extract_image_text(tmp_path) + "\n"
            finally:
                os.unlink(tmp_path)
                
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    
    def _extract_image_text(self, file_path: str) -> str:
        """Enhanced OCR with preprocessing and fallbacks"""
        try:
            # 1) Try EasyOCR on the raw file path
            if easyocr_available:
                try:
                    text = " ".join(ocr_reader.readtext(file_path, detail=0))
                    if text and len(text.strip()) > 10:
                        return text
                except Exception as e:
                    print(f"EasyOCR(file) failed: {e}")

            # 2) Preprocess & try EasyOCR on numpy array
            try:
                preprocessed_img = self.preprocess_image_for_ocr(file_path, resize_scale=1.6, denoise=True, threshold=True)
                if easyocr_available:
                    try:
                        text = " ".join(ocr_reader.readtext(preprocessed_img, detail=0))
                        if text and len(text.strip()) > 8:
                            return text
                    except Exception as e:
                        print(f"EasyOCR(numpy) failed: {e}")
            except Exception as e:
                print(f"Preprocessing failed: {e}")

            # 3) Fallback to Tesseract
            try:
                # Try with preprocessed image first
                if 'preprocessed_img' in locals():
                    text = pytesseract.image_to_string(preprocessed_img)
                else:
                    # Fallback to original image
                    img = Image.open(file_path).convert("RGB")
                    text = pytesseract.image_to_string(img)
                
                if text and len(text.strip()) > 0:
                    return text
            except Exception as e:
                print(f"Tesseract failed: {e}")

            return ""  # Fallback if all methods fail
        except Exception as e:
            raise Exception(f"Image OCR error: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Enhanced text cleaning and normalization"""
        if not text:
            return ""

        # Normalize unicode
        text = unicodedata.normalize("NFKC", text)

        # Common OCR character fixes
        char_map = {
            '“': '"', '”': '"', '‘': "'", '’': "'",
            '`': "'", '´': "'", '″': '"', '„': '"',
            '…': '...', '–': '-', '—': '-', '‒': '-',
        }
        for bad_char, good_char in char_map.items():
            text = text.replace(bad_char, good_char)

        # Fix common OCR confusions
        text = re.sub(r'(?<=\d)[lI](?=\d)', '1', text)   # 1 vs l between digits
        text = re.sub(r'\b0([A-Za-z])', r'O\1', text)    # leading zero before letter -> O
        text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)    # hyphen+newline join
        text = re.sub(r'([a-z])\s+([a-z]{1,2}\b)', r'\1\2', text)  # join broken words

        # Remove control characters and excessive punctuation
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', ' ', text)
        text = re.sub(r'[-]{2,}', '-', text)
        text = re.sub(r'[_]{2,}', ' ', text)
        text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # Keep ASCII only

        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        return text
    
    # The rest of the methods remain the same...
    def extract_basic_info(self, text: str) -> Dict[str, any]:
        """Extract basic information from resume text"""
        return {
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'experience_years': self._extract_experience_years(text),
            'education': self._extract_education_keywords(text),
            'skills': self._extract_basic_skills(text)
        }
    
    def _extract_email(self, text: str) -> Optional[str]:
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        phone_patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}'
        ]
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        return None
    
    def _extract_experience_years(self, text: str) -> int:
        patterns = [
            r'(\d+)\s*\+?\s*years?\s+(?:of\s+)?experience',
            r'(\d+)\s*\+?\s*yrs?\s+(?:of\s+)?experience',
            r'experience\s*:?\s*(\d+)\s*\+?\s*years?',
            r'(\d+)\s*\+?\s*years?\s+in\s+\w+'
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                try:
                    return int(matches[0])
                except ValueError:
                    continue
        return 0
    
    def _extract_education_keywords(self, text: str) -> List[str]:
        keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'degree', 'diploma',
            'engineering', 'computer science', 'mba', 'btech', 'mtech',
            'university', 'college', 'institute', 'bs', 'ms', 'msc', 'bsc'
        ]
        found = [kw for kw in keywords if kw in text.lower()]
        return list(set(found))
    
    def _extract_basic_skills(self, text: str) -> List[str]:
        skill_keywords = [
            'python', 'java', 'javascript', 'typescript', 'react', 'angular',
            'vue.js', 'node.js', 'sql', 'mongodb', 'postgresql', 'mysql',
            'html', 'css', 'bootstrap', 'tailwind', 'git', 'docker',
            'kubernetes', 'aws', 'azure', 'gcp', 'linux',
            'machine learning', 'deep learning', 'data science',
            'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
            'matplotlib', 'seaborn', 'django', 'flask', 'fastapi', 'redis',
            'spring', 'spring boot', 'hadoop', 'spark'
        ]
        return [skill for skill in skill_keywords if skill in text.lower()]